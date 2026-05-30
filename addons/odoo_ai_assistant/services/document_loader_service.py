"""
DocumentLoaderService — 把上傳檔案抽成純文字並切塊（RAG ingestion）。

支援：pdf / docx / xlsx / xls / csv / txt / md
依賴：pypdf（PDF）、python-docx（Word）、openpyxl（Excel，Odoo 內建）
"""
import csv
import io


# ---------------------------------------------------------------------------
# 文字解碼（台灣檔案常見 Big5）
# ---------------------------------------------------------------------------

def _decode(data: bytes) -> str:
    for enc in ('utf-8-sig', 'utf-8', 'big5', 'cp950'):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode('utf-8', errors='ignore')


# ---------------------------------------------------------------------------
# 各格式解析
# ---------------------------------------------------------------------------

# 文字型 PDF 抽出的字數低於此值時，視為掃描型 PDF，改走 OCR
_MIN_PDF_TEXT_LEN = 30


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError('解析 PDF 需要 pypdf：pip install pypdf')
    reader = PdfReader(io.BytesIO(data))
    text = '\n'.join((page.extract_text() or '') for page in reader.pages)
    # 掃描型 PDF（圖片）pypdf 抽不到文字，改用 OCR
    if len(text.strip()) < _MIN_PDF_TEXT_LEN:
        ocr_text = _ocr_pdf(data)
        if ocr_text.strip():
            return ocr_text
    return text


def _ocr_pdf(data: bytes) -> str:
    """用 PyMuPDF 將每頁渲染成圖片，再用 tesseract 做中英文 OCR。"""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError:
        raise RuntimeError(
            '掃描型 PDF 需要 OCR 套件：請確認容器已安裝 '
            'tesseract-ocr（含 chi_tra 語言包）與 pip 套件 pymupdf、pytesseract。'
        )
    doc = fitz.open(stream=data, filetype='pdf')
    texts = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes('png')))
        texts.append(pytesseract.image_to_string(img, lang='chi_tra+chi_sim+eng'))
    return '\n'.join(texts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError('解析 Word 需要 python-docx：pip install python-docx')
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也抽表格內容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError('解析 Excel 需要 openpyxl')
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f'# 工作表：{ws.title}')
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(' | '.join(cells))
    return '\n'.join(lines)


def _extract_csv(data: bytes) -> str:
    text = _decode(data)
    reader = csv.reader(io.StringIO(text))
    return '\n'.join(' | '.join(row) for row in reader if any(row))


# ---------------------------------------------------------------------------
# 對外 API
# ---------------------------------------------------------------------------

def extract_text(data: bytes, filename: str) -> str:
    ext = (filename or '').rsplit('.', 1)[-1].lower() if '.' in (filename or '') else ''
    if ext == 'pdf':
        return _extract_pdf(data)
    if ext == 'docx':
        return _extract_docx(data)
    if ext in ('xlsx', 'xlsm'):
        return _extract_xlsx(data)
    if ext == 'csv':
        return _extract_csv(data)
    if ext in ('txt', 'md', 'markdown', 'text', ''):
        return _decode(data)
    raise RuntimeError(f'尚不支援的檔案格式：.{ext}')


def chunk_text(text: str, chunk_size: int = 600, overlap: int = 100) -> list:
    """
    以字元為單位切塊，塊間保留 overlap 重疊，盡量在換行/句號邊界斷開。
    """
    text = (text or '').strip()
    if not text:
        return []
    chunk_size = max(100, int(chunk_size))
    overlap = max(0, min(int(overlap), chunk_size - 1))

    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        # 嘗試在邊界（換行或句號）斷開，避免切在句中
        if end < n:
            window = text[start:end]
            for sep in ('\n', '。', '！', '？', '. '):
                pos = window.rfind(sep)
                if pos >= chunk_size * 0.5:  # 邊界不能離起點太近
                    end = start + pos + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks
